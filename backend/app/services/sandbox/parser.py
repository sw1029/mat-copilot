from __future__ import annotations

import io
import logging
import re
import stat
import zipfile
from dataclasses import dataclass, field
from pathlib import PurePosixPath

from docx import Document

from app import constants
from app.models.enums import ArtifactIngestStatus
from app.observability import log_sandbox_event

logger = logging.getLogger("app.sandbox.parser")

MAX_ZIP_ENTRIES = getattr(constants, "ZIP_MAX_ENTRIES", 1000)
MAX_ZIP_TOTAL_UNCOMPRESSED = getattr(constants, "ZIP_MAX_TOTAL_BYTES", 100 * 1024 * 1024)
MAX_COMPRESSION_RATIO = 100

TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".go", ".rs",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".rb", ".php", ".swift", ".kt", ".json",
    ".yaml", ".yml", ".toml", ".ini", ".css", ".html", ".csv", ".sql", ".sh",
    ".xml", ".vue", ".svelte", ".scss", ".less", ".dockerfile", ".env", ".cfg",
}


@dataclass
class ParsedFile:
    path: str
    text: str


@dataclass
class SandboxNote:
    path: str
    status: ArtifactIngestStatus
    reason: str


@dataclass
class SandboxResult:
    files: list[ParsedFile] = field(default_factory=list)
    notes: list[SandboxNote] = field(default_factory=list)
    status: ArtifactIngestStatus = ArtifactIngestStatus.SKIPPED_UNSUPPORTED

    @property
    def note_text(self) -> str | None:
        if not self.notes:
            return None
        return "; ".join(f"{n.path}: {n.reason}" for n in self.notes[:8])


def _ext(name: str) -> str:
    return PurePosixPath(name.lower()).suffix


def _safe_zip_path(name: str) -> bool:
    if "\\" in name or re.match(r"^[A-Za-z]:", name):
        return False
    p = PurePosixPath(name)
    return not p.is_absolute() and ".." not in p.parts


def _is_symlink(info: zipfile.ZipInfo) -> bool:
    return stat.S_IFMT(info.external_attr >> 16) == stat.S_IFLNK


def _decode_text(path: str, data: bytes, artifact_id: str) -> ParsedFile | SandboxNote:
    if b"\x00" in data:
        log_sandbox_event(artifact_id=artifact_id, filename=path, size=len(data), result="SKIPPED_UNSUPPORTED", reason="NUL 바이너리")
        return SandboxNote(path, ArtifactIngestStatus.SKIPPED_UNSUPPORTED, "바이너리 파일은 분석에서 제외했습니다.")
    text = data.decode("utf-8", errors="replace")
    log_sandbox_event(artifact_id=artifact_id, filename=path, size=len(data), result="PARSED")
    return ParsedFile(path=path, text=text)


def _parse_docx(path: str, data: bytes, artifact_id: str) -> ParsedFile | SandboxNote:
    try:
        doc = Document(io.BytesIO(data))
        chunks: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append("\t".join(cells))
        text = "\n".join(chunks)
        log_sandbox_event(artifact_id=artifact_id, filename=path, size=len(data), result="PARSED")
        return ParsedFile(path=path, text=text)
    except Exception as exc:  # noqa: BLE001 — 파일별 파싱 실패는 skip note로 표면화
        logger.warning("docx 파싱 실패 path=%s", path, exc_info=True)
        log_sandbox_event(artifact_id=artifact_id, filename=path, size=len(data), result="SKIPPED_UNSUPPORTED", reason=str(exc)[:120])
        return SandboxNote(path, ArtifactIngestStatus.SKIPPED_UNSUPPORTED, "docx 파싱에 실패했습니다.")


def _parse_single(path: str, data: bytes, artifact_id: str) -> ParsedFile | SandboxNote:
    ext = _ext(path)
    if ext == ".docx":
        return _parse_docx(path, data, artifact_id)
    if ext in TEXT_EXTENSIONS or ext == "":
        return _decode_text(path, data, artifact_id)
    log_sandbox_event(artifact_id=artifact_id, filename=path, size=len(data), result="SKIPPED_UNSUPPORTED", reason="unsupported extension")
    return SandboxNote(path, ArtifactIngestStatus.SKIPPED_UNSUPPORTED, "지원하지 않는 파일 형식입니다.")


def parse_upload(filename: str, data: bytes, *, artifact_id: str) -> SandboxResult:
    """TRD §8 sandbox: 실행·네트워크 없이 허용 포맷만 정적 텍스트 추출."""
    if _ext(filename) == ".zip":
        return _parse_zip(filename, data, artifact_id=artifact_id)
    item = _parse_single(filename, data, artifact_id)
    result = SandboxResult()
    if isinstance(item, ParsedFile):
        result.files.append(item)
        result.status = ArtifactIngestStatus.PARSED
    else:
        result.notes.append(item)
        result.status = item.status
    return result


def _parse_zip(filename: str, data: bytes, *, artifact_id: str) -> SandboxResult:
    result = SandboxResult()
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile:
        note = SandboxNote(filename, ArtifactIngestStatus.SKIPPED_UNSUPPORTED, "zip 파일을 열 수 없습니다.")
        result.notes.append(note)
        result.status = note.status
        return result

    infos = zf.infolist()
    if len(infos) > MAX_ZIP_ENTRIES:
        result.notes.append(SandboxNote(filename, ArtifactIngestStatus.SKIPPED_TOO_LARGE, "zip 엔트리 수 한도를 초과했습니다."))
        result.status = ArtifactIngestStatus.SKIPPED_TOO_LARGE
        log_sandbox_event(artifact_id=artifact_id, filename=filename, size=len(data), result="SKIPPED_TOO_LARGE", reason="too many entries")
        return result

    total = 0
    for info in infos:
        if not _safe_zip_path(info.filename):
            result.notes.append(SandboxNote(info.filename, ArtifactIngestStatus.BLOCKED_UNSAFE, "zip 경로 탈출이 감지되어 전체 파일을 차단했습니다."))
            result.status = ArtifactIngestStatus.BLOCKED_UNSAFE
            log_sandbox_event(artifact_id=artifact_id, filename=info.filename, size=info.file_size, result="BLOCKED_UNSAFE", reason="path traversal")
            return result
        total += info.file_size
        if total > MAX_ZIP_TOTAL_UNCOMPRESSED:
            result.notes.append(SandboxNote(filename, ArtifactIngestStatus.SKIPPED_TOO_LARGE, "zip 해제 총량 한도를 초과했습니다."))
            result.status = ArtifactIngestStatus.SKIPPED_TOO_LARGE
            log_sandbox_event(artifact_id=artifact_id, filename=filename, size=total, result="SKIPPED_TOO_LARGE", reason="uncompressed total")
            return result
        if info.compress_size and info.file_size / max(info.compress_size, 1) > MAX_COMPRESSION_RATIO:
            result.notes.append(SandboxNote(info.filename, ArtifactIngestStatus.SKIPPED_TOO_LARGE, "압축비가 비정상적으로 높아 zip 폭탄으로 간주했습니다."))
            result.status = ArtifactIngestStatus.SKIPPED_TOO_LARGE
            log_sandbox_event(artifact_id=artifact_id, filename=info.filename, size=info.file_size, result="SKIPPED_TOO_LARGE", reason="compression ratio")
            return result

    for info in infos:
        if info.is_dir():
            continue
        path = info.filename
        if _is_symlink(info):
            note = SandboxNote(path, ArtifactIngestStatus.SKIPPED_UNSUPPORTED, "심볼릭 링크 엔트리는 제외했습니다.")
            result.notes.append(note)
            log_sandbox_event(artifact_id=artifact_id, filename=path, size=0, result="SKIPPED_UNSUPPORTED", reason="symlink")
            continue
        if _ext(path) == ".zip":
            note = SandboxNote(path, ArtifactIngestStatus.SKIPPED_UNSUPPORTED, "중첩 zip은 해제하지 않았습니다.")
            result.notes.append(note)
            log_sandbox_event(artifact_id=artifact_id, filename=path, size=info.file_size, result="SKIPPED_UNSUPPORTED", reason="nested zip")
            continue
        item = _parse_single(path, zf.read(info), artifact_id)
        if isinstance(item, ParsedFile):
            result.files.append(item)
        else:
            result.notes.append(item)

    if result.files:
        result.status = ArtifactIngestStatus.PARSED
    elif result.notes:
        if any(n.status is ArtifactIngestStatus.BLOCKED_UNSAFE for n in result.notes):
            result.status = ArtifactIngestStatus.BLOCKED_UNSAFE
        elif any(n.status is ArtifactIngestStatus.SKIPPED_TOO_LARGE for n in result.notes):
            result.status = ArtifactIngestStatus.SKIPPED_TOO_LARGE
        else:
            result.status = ArtifactIngestStatus.SKIPPED_UNSUPPORTED
    else:
        result.status = ArtifactIngestStatus.SKIPPED_UNSUPPORTED
    return result


def whitespace_normalized_contains(haystack: str, needle: str) -> bool:
    norm = lambda s: re.sub(r"\s+", " ", s).strip()
    return norm(needle) in norm(haystack)
